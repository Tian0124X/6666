"""数据分析工具 — Excel/CSV → 清洗 → 统计 → 图表 → Word 报告"""

import os
import logging
from typing import Optional, Literal
from pydantic import BaseModel, Field
from langchain_core.tools import BaseTool
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from app.tools.base import register_tool
from app.tools.registry import validate_file_path, ensure_directory

logger = logging.getLogger(__name__)

# 中文字体支持
plt.rcParams["font.sans-serif"] = ["SimHei", "Microsoft YaHei", "DejaVu Sans"]
plt.rcParams["axes.unicode_minus"] = False


class DataAnalyzerInput(BaseModel):
    """入参 schema（Pydantic 严格校验）"""
    file_path: str = Field(description="Excel/CSV 文件的绝对路径")
    action: Literal["summary", "analyze", "full_report"] = Field(
        default="summary",
        description="summary=概览统计 | analyze=深度分析+图表 | full_report=生成 Word 报告",
    )
    target_column: Optional[str] = Field(default=None, description="重点分析的列名")
    chart_type: Optional[Literal["bar", "line", "pie", "scatter"]] = Field(
        default=None, description="图表类型"
    )


@register_tool
class DataAnalyzerTool(BaseTool):
    """数据分析工具：读文件 → 清洗 → 统计 → 图表 → Word 报告"""
    name: str = "data_analyzer"
    description: str = (
        "分析 Excel/CSV 数据文件。支持 summary(概览)、analyze(深度分析+图表)、"
        "full_report(生成含图表的 Word 报告)。可指定列名和图表类型(bar/line/pie/scatter)。"
    )
    args_schema: type[BaseModel] = DataAnalyzerInput

    # ====== 内部方法 ======

    def _load_data(self, file_path: str) -> pd.DataFrame:
        """加载 Excel/CSV，自动检测编码"""
        ext = os.path.splitext(file_path)[1].lower()
        if ext in (".xlsx", ".xls"):
            return pd.read_excel(file_path)
        elif ext == ".csv":
            for enc in ["utf-8", "gbk", "gb2312", "latin-1"]:
                try:
                    return pd.read_csv(file_path, encoding=enc)
                except UnicodeDecodeError:
                    continue
            raise ValueError("无法识别 CSV 编码，已尝试 utf-8/gbk/gb2312/latin-1")
        else:
            raise ValueError(f"不支持的文件格式: {ext}，仅支持 xlsx/xls/csv")

    def _clean_data(self, df: pd.DataFrame) -> pd.DataFrame:
        """清洗：去重 + 中位数填充数值列"""
        df = df.copy()
        before = len(df)
        df = df.drop_duplicates()
        removed = before - len(df)
        if removed > 0:
            logger.info(f"去重: {before} → {len(df)} (移除 {removed} 行)")

        numeric_cols = df.select_dtypes(include=["number"]).columns
        for col in numeric_cols:
            missing = df[col].isna().sum()
            if missing > 0:
                df[col] = df[col].fillna(df[col].median())
                logger.info(f"中位数填充: {col} ({missing} 个缺失值)")
        return df

    def _generate_summary(self, df: pd.DataFrame) -> str:
        """生成描述统计"""
        lines = [
            f"## 数据概览",
            f"- 行数: {len(df)} | 列数: {len(df.columns)}",
            f"- 列名: {', '.join(df.columns.tolist())}",
            f"- 缺失值: {df.isna().sum().sum()}",
            f"",
            f"### 数值列统计",
            f"```",
            df.describe().to_string(),
            f"```",
        ]
        return "\n".join(lines)

    def _generate_chart(
        self, df: pd.DataFrame, target_column: Optional[str],
        chart_type: str = "bar", output_dir: str = "data/reports",
    ) -> Optional[str]:
        """生成 matplotlib 图表，返回图片路径"""
        ensure_directory(output_dir)

        numeric_cols = df.select_dtypes(include=["number"]).columns
        if len(numeric_cols) == 0:
            logger.warning("无数值列，跳过图表")
            return None

        col = target_column or numeric_cols[0]
        if col not in df.columns:
            col = numeric_cols[0]

        # 安全清理列名和图表类型，防止路径遍历
        import re
        safe_col = re.sub(r'[\\/:*?"<>|.]', '_', str(col))
        safe_type = re.sub(r'[\\/:*?"<>|.]', '_', str(chart_type))

        fig, ax = plt.subplots(figsize=(10, 6))

        if chart_type == "bar":
            data = df[col].value_counts().head(20)
            ax.bar(range(len(data)), data.values)
            ax.set_xticks(range(len(data)))
            ax.set_xticklabels(data.index, rotation=45, ha="right")
            ax.set_title(f"{col} 分布 (柱状图)")
        elif chart_type == "line":
            data = df[col].head(50)
            ax.plot(range(len(data)), data.values, marker="o", markersize=3)
            ax.set_title(f"{col} 趋势 (折线图)")
        elif chart_type == "pie":
            data = df[col].value_counts().head(10)
            ax.pie(data.values, labels=data.index, autopct="%1.1f%%")
            ax.set_title(f"{col} 占比 (饼图)")
        else:  # scatter
            if len(numeric_cols) >= 2:
                ax.scatter(df[numeric_cols[0]], df[numeric_cols[1]], alpha=0.5)
                ax.set_xlabel(numeric_cols[0])
                ax.set_ylabel(numeric_cols[1])
                ax.set_title(f"{numeric_cols[0]} vs {numeric_cols[1]} (散点图)")
            else:
                return None

        chart_path = os.path.join(output_dir, f"chart_{safe_col}_{safe_type}.png")
        plt.tight_layout()
        fig.savefig(chart_path, dpi=150)
        plt.close(fig)
        logger.info(f"图表已保存: {chart_path}")
        return chart_path

    def _generate_word_report(
        self, df: pd.DataFrame, summary: str,
        chart_paths: list[str], output_dir: str = "data/reports",
    ) -> str:
        """生成 Word 报告（含图表）"""
        from docx import Document
        from docx.shared import Inches, Pt
        ensure_directory(output_dir)

        doc = Document()
        doc.add_heading("数据分析报告", level=0)
        doc.add_paragraph(f"生成时间: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}")

        doc.add_heading("数据概览", level=1)
        doc.add_paragraph(f"行数: {len(df)} | 列数: {len(df.columns)}")
        doc.add_paragraph(f"列名: {', '.join(df.columns.tolist())}")

        doc.add_heading("描述统计", level=1)
        for line in str(df.describe()).split("\n"):
            doc.add_paragraph(line)

        if chart_paths:
            doc.add_heading("可视化图表", level=1)
            for path in chart_paths:
                if os.path.exists(path):
                    doc.add_picture(path, width=Inches(5.5))
                    caption = os.path.basename(path).replace(".png", "")
                    doc.add_paragraph(caption)

        report_path = os.path.join(
            output_dir,
            f"report_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.docx",
        )
        doc.save(report_path)
        return report_path

    # ====== 主入口 ======

    def _run(
        self,
        file_path: str,
        action: str = "summary",
        target_column: Optional[str] = None,
        chart_type: Optional[str] = None,
    ) -> str:
        try:
            # 安全校验
            safe_path = validate_file_path(file_path)
            if not os.path.exists(safe_path):
                return f"❌ 文件不存在: {safe_path}"

            # 加载 + 清洗
            df = self._load_data(safe_path)
            df = self._clean_data(df)

            if action == "summary":
                return self._generate_summary(df)

            elif action == "analyze":
                summary = self._generate_summary(df)
                chart_path = None
                if chart_type:
                    chart_path = self._generate_chart(df, target_column, chart_type)
                result = summary
                if chart_path:
                    result += f"\n\n📊 图表: {chart_path}"
                return result

            elif action == "full_report":
                summary = self._generate_summary(df)
                chart_paths = []
                for ct in (["bar", "line"] if not chart_type else [chart_type]):
                    p = self._generate_chart(df, target_column, ct)
                    if p:
                        chart_paths.append(p)
                report_path = self._generate_word_report(df, summary, chart_paths)
                return (
                    f"✅ 报告已生成: {report_path}\n\n"
                    f"{summary}\n\n"
                    f"📊 图表数量: {len(chart_paths)}"
                )

        except ValueError as e:
            return f"❌ 参数错误: {e}"
        except Exception as e:
            logger.error(f"数据分析失败: {e}", exc_info=True)
            return f"❌ 分析失败: {e}"
