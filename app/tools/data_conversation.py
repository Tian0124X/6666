"""
数据对话引擎 — LLM 驱动的自然语言数据分析

上传 Excel/CSV → 自然语言提问 → LLM 生成 pandas 代码 → 安全执行 → 图表+解读
"""

import os
import re
import json
import time
import logging
import traceback
from typing import Optional, Any
from functools import lru_cache
from pydantic import BaseModel, Field, model_validator
from langchain_core.tools import BaseTool
from langchain_openai import ChatOpenAI
from langchain_core.messages import SystemMessage, HumanMessage
import pandas as pd
import numpy as np

from app.config import settings
from app.tools.base import register_tool
from app.tools.registry import validate_file_path

logger = logging.getLogger(__name__)

# ====== 安全沙箱 ======

SAFE_BUILTINS = {
    "len": len, "range": range, "list": list, "dict": dict, "tuple": tuple,
    "str": str, "int": int, "float": float, "bool": bool, "complex": complex,
    "print": print, "zip": zip, "enumerate": enumerate,
    "sorted": sorted, "min": min, "max": max, "sum": sum, "abs": abs, "round": round,
    "set": set, "frozenset": frozenset, "reversed": reversed,
    "True": True, "False": False, "None": None,
    "type": type, "isinstance": isinstance, "hasattr": hasattr, "getattr": getattr,
    "map": map, "filter": filter, "any": any, "all": all,
    "format": format, "divmod": divmod, "pow": pow,
    "slice": slice, "KeyError": KeyError, "ValueError": ValueError,
    "TypeError": TypeError, "IndexError": IndexError,
}

FORBIDDEN_PATTERNS = [
    r'__', r'import\s', r'from\s+\w+\s+import',
    r'open\s*\(', r'eval\s*\(', r'exec\s*\(',
    r'os\.', r'sys\.', r'subprocess', r'shutil',
    r'globals\s*\(', r'locals\s*\(', r'getattr\s*\(.*__',
    r'compile\s*\(', r'breakpoint\s*\(', r'input\s*\(',
    r'\._', r'\.__', r'ctypes', r'builtins',
    r'write\s*\(', r'\.save\s*\(', r'\.to_excel', r'\.to_csv',
    r'requests?\.', r'urllib', r'http', r'socket',
    r'multiprocessing', r'threading', r'signal',
    r'os\.system', r'os\.popen', r'os\.spawn',
    r'matplotlib', r'plt\.', r'base64', r'BytesIO', r'savefig',
    r'pyplot', r'figure\s*\(', r'imshow', r'imread',
]


def _sanitize_code(code: str) -> str:
    """检查代码安全性，抛出异常如果不安全"""
    code_lower = code.lower()
    for pattern in FORBIDDEN_PATTERNS:
        if re.search(pattern, code_lower):
            raise ValueError(f"代码包含禁止的操作: {pattern}")
    if re.search(r'__\w+', code) or re.search(r'\w+__', code):
        raise ValueError("代码包含禁止的双下划线访问")
    return code


def _execute_sandbox(code: str, df: pd.DataFrame) -> dict:
    """
    安全沙箱执行 pandas 代码。

    Returns:
        {"result": any, "type": "dataframe"|"series"|"scalar"|"error", "error": str|None}
    """
    _sanitize_code(code)

    namespace = {
        "__builtins__": SAFE_BUILTINS,
        "pd": pd,
        "np": np,
        "df": df,
    }

    try:
        lines = code.strip().split("\n")
        last_line = lines[-1].strip()

        if not last_line.startswith(("if ", "for ", "while ", "def ", "class ", "try:", "except", "else:", "elif ", "with ")) and "=" not in last_line.split("(")[0]:
            exec_lines = "\n".join(lines[:-1]) if len(lines) > 1 else ""
            exec_lines += f"\n__result__ = {last_line}"
        else:
            exec_lines = code
            exec_lines += "\n__result__ = None"

        exec(exec_lines, namespace)

        result = namespace.get("__result__", None)

        if isinstance(result, pd.DataFrame):
            if len(result) > 100:
                result = result.head(100)
            return {
                "type": "dataframe",
                "columns": result.columns.tolist(),
                "rows": result.fillna("").values.tolist()[:100],
                "shape": [len(result), len(result.columns)],
            }
        elif isinstance(result, pd.Series):
            data = result.head(50).to_dict()
            return {"type": "series", "data": data, "name": str(result.name)}
        elif isinstance(result, (int, float, np.integer, np.floating)):
            return {"type": "scalar", "value": float(result) if isinstance(result, (float, np.floating)) else int(result)}
        elif isinstance(result, str):
            return {"type": "scalar", "value": result}
        elif result is None:
            return {"type": "scalar", "value": "执行完成 (无返回值)"}
        else:
            return {"type": "scalar", "value": str(result)}

    except Exception as e:
        return {"type": "error", "error": f"{type(e).__name__}: {e}"}


# ====== 结构化输出模型 ======

class ChartConfig(BaseModel):
    """图表配置 — LLM 结构化输出的图表部分"""
    type: str = Field(default="bar", description="图表类型: bar / line / pie / area / scatter / funnel / composed")
    x: str = Field(default="", description="X轴 / 分类列名")
    y: str = Field(default="", description="Y轴 / 数值列名")
    x2: str = Field(default="", description="第二数值列(散点图专用)")
    title: str = Field(default="图表", description="图表标题")
    data: list[dict[str, Any]] = Field(default_factory=list, description="图表数据，最多20条")
    series: list[dict[str, Any]] | None = Field(default=None, description="多系列数据(组合图)")

    @model_validator(mode="before")
    @classmethod
    def _fix_list_y(cls, values: Any) -> Any:
        """兼容 LLM 将 y 输出为列表的情况（多指标场景）"""
        if not isinstance(values, dict):
            return values
        y = values.get("y")
        if isinstance(y, list):
            # y 是列表 → 取第一个作为 y，剩余生成 series
            if y:
                values["y"] = str(y[0])
                if len(y) > 1 and not values.get("series"):
                    values["series"] = [
                        {"dataKey": str(col), "chartType": "bar" if i == 0 else "line"}
                        for i, col in enumerate(y)
                    ]
                if values.get("type") not in ("composed",):
                    if len(y) > 1:
                        values["type"] = "composed"
        return values


class DataAnalysisOutput(BaseModel):
    """LLM 输出 — 数据分析结果（灵活版）"""
    content: str = Field(description="markdown 格式的自然语言回答，这是用户看到的主要内容")
    code: str = Field(default="", description="可选的 pandas 代码，直接操作 df，最后一行是表达式")
    chart_config: ChartConfig | None = Field(default=None, description="可选图表配置，仅当图表能增强理解时输出")
    suggested_questions: list[str] = Field(default_factory=list, description="推荐的后续分析问题(2-3个)")


# ====== LLM 系统提示 ======

DATA_CHAT_SYSTEM = """你是一位资深企业数据分析师。用户上传了数据文件，已加载为 df 变量。

## 数据概况
{df_info}

## 核心原则

### 1. 意图优先
先判断用户的真实意图：
- 用户只是打招呼/闲聊 → 自然回应，不要硬分析数据
- 用户问"数据怎么样" → 给概览统计 + 关键发现
- 用户问具体问题 → 精准回答 + 数据支撑
- 用户说"生成报告" → 标记需要报告输出

### 2. 对话自然 — 像人类数据分析师
- 先给结论和洞察，再附细节
- 主动发现数据中的异常值、趋势拐点、分布特征
- 给出业务建议和可行操作方向，而不仅仅是呈现数据
- **重要：content 中不要写 markdown 表格**（表格数据会由系统自动渲染在下方）

### 3. 按需输出 — 根据用户问题灵活决定
- 概览型问题（"数据怎么样""有几列"）→ 纯文字回答，不需要代码和图表
- 排名/对比问题（"最高""最低""Top10"）→ 文字结论 + 表格 + 柱状图
- 占比/分布问题（"占比""分布""构成"）→ 文字结论 + 饼图/面积图
- 趋势问题（"趋势""变化""增长"）→ 文字结论 + 折线图/面积图
- 相关性问题（"关系""关联""影响"）→ 文字结论 + 散点图
- 漏斗/转化问题（"转化率""漏斗""流程"）→ 文字结论 + 漏斗图
- 多维度问题（"同时看""对比趋势"）→ 文字结论 + 组合图
- 纯计算问题（"平均""总和"）→ 文字结论即可，不需要图表

### 4. 图表智能选择（仅在"一图胜千言"时输出 chart_config）
可选类型: bar(柱状图) / line(折线图) / pie(饼图) / area(面积图) / scatter(散点图) / funnel(漏斗图) / composed(组合图)
- 时间序列 + 累计值 → type: "area"（面积图，展示趋势填充效果更佳）
- 纯时间序列趋势 → type: "line"（折线图）
- 分类占比(≤6个) → type: "pie"（饼图）
- 分类对比(7~20个) → type: "bar"（柱状图）
- 两个数值列相关性 → type: "scatter"（散点图），设置 x2 为第二列
- 有序递减序列 → type: "funnel"（漏斗图）
- 同时展示数值+趋势 → type: "composed"（组合图=柱+线），提供 series 字段
- 类别>20个 → 取Top10做柱状图
- 纯统计数字 → 不需要图表

### 5. 代码原则
- 直接操作 df 变量，禁止 import / read_excel / read_csv / open / print()
- 最后一行是表达式（其值会被自动捕获为结果）
- **查询 Top N 时使用 head(10) 或 nlargest(10)，不要只取 head(3)**
- 如果不需要计算就不要写 code（设为空字符串）

### 6. 后续建议
回答末尾通过 suggested_questions 推荐 2-3 个后续分析方向，引导用户深入探索数据。

## 输出 JSON 格式
{{
  "content": "markdown 格式的自然语言回答（主要输出）",
  "code": "可选的 pandas 代码（不需要时为空字符串）",
  "chart_config": {{"type":"bar|line|pie|area|scatter|funnel|composed","x":"列名","y":"列名","title":"标题","data":[{{...}}]}} 或 null,
  "suggested_questions": ["后续分析建议1", "后续分析建议2"]
}}

注意: chart_config 和 code 仅在必要时输出。不需要就设为 null / ""。
"""


def _build_df_info(df: pd.DataFrame) -> str:
    """构建 DataFrame 的摘要信息给 LLM"""
    lines = [f"- 行数: {len(df)}, 列数: {len(df.columns)}"]
    lines.append("- 列名及类型:")
    for col in df.columns:
        dtype = str(df[col].dtype)
        sample = df[col].dropna().head(3).tolist()
        sample_str = ", ".join(str(x)[:50] for x in sample)
        nulls = df[col].isna().sum()
        info = f"  {col} ({dtype})"
        if sample_str:
            info += f" 样例: [{sample_str}]"
        if nulls > 0:
            info += f" 缺失: {nulls}"
        lines.append(info)
    return "\n".join(lines)


def _extract_json(text: str) -> dict:
    """
    从 LLM 响应中提取 JSON 对象。
    处理 markdown 代码块包裹、首尾多余文字等情况。
    """
    # 尝试直接解析
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # 提取 markdown ```json ... ``` 代码块
    m = re.search(r'```(?:json)?\s*\n?(.*?)\n?```', text, re.DOTALL)
    if m:
        try:
            return json.loads(m.group(1).strip())
        except json.JSONDecodeError:
            pass

    # 提取 { ... } 最外层 JSON 对象
    m = re.search(r'\{.*\}', text, re.DOTALL)
    if m:
        try:
            return json.loads(m.group(0))
        except json.JSONDecodeError:
            pass

    raise ValueError(f"无法从响应中提取 JSON: {text[:200]}...")


def _get_llm():
    """获取 LLM 实例"""
    return ChatOpenAI(
        model=settings.LLM_MODEL,
        api_key=settings.LLM_API_KEY,
        base_url=settings.LLM_BASE_URL,
        temperature=0,
        timeout=settings.LLM_TIMEOUT,
        max_tokens=4096,
    )


def _is_datetime_column(series: pd.Series) -> bool:
    """检测列是否为时间/日期类型"""
    if pd.api.types.is_datetime64_any_dtype(series):
        return True
    if series.dtype == 'object':
        try:
            pd.to_datetime(series.dropna().head(10))
            return True
        except (ValueError, TypeError):
            pass
    return False


# ====== DataFrame 缓存 — 消除重复加载 ======

_dataframe_cache: dict[str, tuple[float, pd.DataFrame]] = {}
_CACHE_MAX_SIZE = 8


def _get_dataframe(file_path: str) -> pd.DataFrame:
    """
    带缓存的 DataFrame 加载。
    按文件路径+修改时间做 key，避免每次对话都重新读取文件。
    缓存上限 8 个文件，LRU 淘汰。
    """
    mtime = os.path.getmtime(file_path)
    cached = _dataframe_cache.get(file_path)
    if cached and cached[0] == mtime:
        logger.debug(f"DataFrame 缓存命中: {file_path}")
        return cached[1].copy()

    ext = os.path.splitext(file_path)[1].lower()
    if ext in (".xlsx", ".xls"):
        df = pd.read_excel(file_path)
        unnamed_count = sum(1 for c in df.columns if 'Unnamed' in str(c))
        if unnamed_count > len(df.columns) * 0.5:
            logger.info(f"检测到标题行，使用 header=1 ({unnamed_count}/{len(df.columns)} Unnamed)")
            df = pd.read_excel(file_path, header=1)
        df.columns = [str(c).strip() for c in df.columns]
        df = df.dropna(axis=1, how='all')
    elif ext == ".csv":
        for enc in ["utf-8", "gbk", "gb2312", "latin-1", "utf-16"]:
            try:
                df = pd.read_csv(file_path, encoding=enc)
                break
            except (UnicodeDecodeError, UnicodeError):
                continue
        else:
            raise ValueError("无法识别 CSV 编码")
    else:
        raise ValueError(f"不支持的文件格式: {ext}")

    # LRU 淘汰
    if len(_dataframe_cache) >= _CACHE_MAX_SIZE:
        oldest = sorted(_dataframe_cache.items(), key=lambda x: x[1][0])
        for k, _ in oldest[:len(oldest) // 2]:
            del _dataframe_cache[k]

    _dataframe_cache[file_path] = (mtime, df)
    logger.info(f"DataFrame 已加载并缓存: {file_path} ({len(df)}x{len(df.columns)})")
    return df.copy()


# ====== 数据洞察引擎 ======

def _generate_data_insights(df: pd.DataFrame, result: dict | None = None) -> dict:
    """
    自动数据洞察引擎。
    生成数据摘要、异常值检测、相关性分析、分布特征。
    """
    insights: dict[str, Any] = {
        "summary": "",
        "anomalies": [],
        "correlations": [],
        "suggestions": [],
    }

    n_rows, n_cols = len(df), len(df.columns)
    numeric_cols = df.select_dtypes(include=["number"]).columns.tolist()
    cat_cols = df.select_dtypes(include=["object", "category"]).columns.tolist()
    total_missing = int(df.isna().sum().sum())

    # 1. 数据摘要
    summary_parts = [f"数据集包含 {n_rows} 行 x {n_cols} 列"]
    if numeric_cols:
        summary_parts.append(f"{len(numeric_cols)} 个数值列")
    if cat_cols:
        summary_parts.append(f"{len(cat_cols)} 个分类列")
    if total_missing > 0:
        missing_pct = total_missing / (n_rows * n_cols) * 100
        summary_parts.append(f"缺失值 {total_missing} 个 ({missing_pct:.1f}%)")
    else:
        summary_parts.append("无缺失值")
    insights["summary"] = "，".join(summary_parts) + "。"

    # 2. 异常值检测 (IQR 方法，最多检查 5 个数值列)
    for col in numeric_cols[:5]:
        q1 = df[col].quantile(0.25)
        q3 = df[col].quantile(0.75)
        iqr = q3 - q1
        lower = q1 - 1.5 * iqr
        upper = q3 + 1.5 * iqr
        n_outliers = int(((df[col] < lower) | (df[col] > upper)).sum())
        if n_outliers > 0:
            pct = n_outliers / n_rows * 100
            insights["anomalies"].append({
                "column": col,
                "count": n_outliers,
                "percentage": round(pct, 1),
                "range": f"[{lower:.2f}, {upper:.2f}]",
            })

    # 3. 相关性分析 (数值列间，最多 6 列)
    if len(numeric_cols) >= 2:
        corr_cols = numeric_cols[:6]
        corr_matrix = df[corr_cols].corr()
        for i in range(len(corr_cols)):
            for j in range(i + 1, len(corr_cols)):
                r = corr_matrix.iloc[i, j]
                if abs(r) > 0.7 and not pd.isna(r):
                    direction = "正相关" if r > 0 else "负相关"
                    strength = "强" if abs(r) > 0.85 else "较强"
                    insights["correlations"].append({
                        "col_a": corr_cols[i],
                        "col_b": corr_cols[j],
                        "value": round(float(r), 3),
                        "description": f"{corr_cols[i]} 与 {corr_cols[j]} {strength}{direction}(r={r:.2f})",
                    })

    # 4. 分布特征 (前 3 个数值列)
    for col in numeric_cols[:3]:
        skew = df[col].skew()
        if abs(skew) > 1:
            direction = "右偏" if skew > 0 else "左偏"
            insights["anomalies"].append({
                "column": col,
                "count": 0,
                "percentage": 0,
                "range": "",
                "description": f"{col} 分布明显{direction}(偏度={skew:.2f})",
            })

    # 5. 分析建议
    suggestions = []
    if len(cat_cols) > 0 and len(numeric_cols) > 0:
        suggestions.append(f"尝试按 {cat_cols[0]} 分组分析 {numeric_cols[0]}")
    if len(numeric_cols) >= 2:
        suggestions.append(f"探索 {numeric_cols[0]} 和 {numeric_cols[1]} 的关系")
    if any(_is_datetime_column(df[c]) for c in df.columns if c in df.columns):
        suggestions.append("按时间维度分析趋势变化")
    if insights["anomalies"]:
        suggestions.append("深入调查检测到的异常值")
    insights["suggestions"] = suggestions[:3]

    return insights


def _sanitize_chart_data(chart: dict | None) -> dict | None:
    """清理图表数据: NaN→0, None→0(数值列)或''(字符串列), 确保key一致"""
    if not chart or not chart.get("data"):
        return chart
    data = chart["data"]
    if not isinstance(data, list) or not data:
        return chart
    for row in data:
        if not isinstance(row, dict):
            continue
        for k, v in row.items():
            if v is None:
                row[k] = 0
            elif isinstance(v, float) and (v != v):  # NaN check
                row[k] = 0
    return chart


def _smart_chart_type(
    df: pd.DataFrame,
    result: dict,
    user_question: str = "",
) -> dict | None:
    """
    基于数据特征智能选择图表类型。
    返回图表配置 dict 或 None（不需要图表）。

    判断逻辑:
    1. 时间序列 + 累计值 → area
    2. 时间序列 → line
    3. 低基数分类(≤6) → pie
    4. 中基数分类(7-20) → bar
    5. 高基数(>20) → bar (Top10)
    6. 两数值列相关性 → scatter
    7. 有序递减序列 → funnel
    8. 多指标对比 → composed
    9. 纯标量结果 → None (不需要图表)
    """
    # 标量结果 → 不需要图表
    if result.get("type") == "scalar":
        return None

    if result.get("type") == "dataframe":
        df_result = pd.DataFrame(
            result["rows"], columns=result["columns"]
        )
        numeric_cols = df_result.select_dtypes(include=["number"]).columns.tolist()
        if not numeric_cols or len(df_result.columns) < 2:
            return None

        x_col = df_result.columns[0]
        y_col = numeric_cols[0]
        n_categories = len(df_result)

        chart_data = [
            dict(zip(df_result.columns.tolist(), row))
            for row in result["rows"][:20]
        ]

        # 1. 时间序列 + 累计值 → area 面积图
        if x_col in df.columns and _is_datetime_column(df[x_col]):
            # 检测是否为累计值(单调递增)
            vals = df_result[y_col].values
            is_cumulative = len(vals) >= 3 and all(
                vals[i] >= vals[i-1] * 0.9 for i in range(1, len(vals))
            )
            if is_cumulative:
                return {
                    "type": "area",
                    "x": x_col, "y": y_col,
                    "title": f"{y_col} 累计趋势",
                    "data": chart_data,
                }
            # 普通时间序列 → line
            return {
                "type": "line",
                "x": x_col, "y": y_col,
                "title": f"{y_col} 变化趋势",
                "data": chart_data,
            }

        # 2. 多数值列 → composed 组合图 (柱+线)
        if len(numeric_cols) >= 3 and n_categories <= 12:
            return {
                "type": "composed",
                "x": x_col, "y": y_col,
                "title": f"{y_col} 与 {numeric_cols[1]} 对比",
                "data": chart_data,
                "series": [
                    {"dataKey": y_col, "chartType": "bar"},
                    {"dataKey": numeric_cols[1], "chartType": "line"},
                ],
            }

        # 3. 有序递减序列 → funnel 漏斗图
        if n_categories >= 3 and n_categories <= 8:
            vals = df_result[y_col].values
            is_decreasing = all(vals[i] >= vals[i+1] * 0.8 for i in range(len(vals)-1))
            if is_decreasing:
                return {
                    "type": "funnel",
                    "x": x_col, "y": y_col,
                    "title": f"{y_col} 漏斗分析",
                    "data": chart_data,
                }

        # 4. 低基数 → 饼图
        if n_categories <= 6:
            return {
                "type": "pie", "x": x_col, "y": y_col,
                "title": f"{y_col} 按{x_col}分布",
                "data": chart_data,
            }

        # 5. 两数值列 → scatter 散点图
        if len(numeric_cols) >= 2:
            q_lower = r"(关系|关联|相关|影响|scatter|散点)"
            if re.search(q_lower, user_question, re.IGNORECASE):
                return {
                    "type": "scatter",
                    "x": numeric_cols[0], "y": numeric_cols[1],
                    "x2": numeric_cols[1],
                    "title": f"{numeric_cols[0]} vs {numeric_cols[1]}",
                    "data": chart_data,
                }

        # 6. 中/高基数 → 柱状图
        return {
            "type": "bar", "x": x_col, "y": y_col,
            "title": f"{y_col} 按{x_col}排名",
            "data": chart_data,
        }

    if result.get("type") == "series":
        data = result.get("data", {})
        if data and len(data) >= 1:
            series_name = result.get("name", "值")
            n_cat = len(data)
            keys = list(data.keys())

            # 检测 key 是否为日期
            is_date_keys = False
            try:
                pd.to_datetime(pd.Series(keys[:10]), format='mixed', dayfirst=False)
                is_date_keys = True
            except (ValueError, TypeError):
                pass

            series_data = [
                {"类别": str(k), series_name: v}
                for k, v in list(data.items())[:20]
            ]

            if is_date_keys and n_cat >= 3:
                # 时间序列: 检查是否累计
                vals = list(data.values())[:20]
                is_cumulative = all(vals[i] >= vals[i-1] * 0.9 for i in range(1, len(vals)))
                chart_type = "area" if is_cumulative else "line"
                title = f"{series_name} 累计趋势" if is_cumulative else f"{series_name} 变化趋势"
                return {
                    "type": chart_type, "x": "类别", "y": series_name,
                    "title": title, "data": series_data,
                }
            elif n_cat <= 6:
                return {
                    "type": "pie", "x": "类别", "y": series_name,
                    "title": f"{series_name} 分布",
                    "data": series_data,
                }
            else:
                return {
                    "type": "bar", "x": "类别", "y": series_name,
                    "title": f"{series_name} 排名 (Top 20)",
                    "data": series_data,
                }

    return None


def _build_chart_from_result(result: dict) -> dict:
    """从代码执行结果自动构建图表配置（向后兼容包装）"""
    return _smart_chart_type(pd.DataFrame(), result) or {}


def analyze_with_llm(file_path: str, question: str, with_chart: bool = True) -> dict:
    """
    主函数: LLM 驱动的数据分析对话 - 使用结构化输出。

    Args:
        file_path: Excel/CSV 文件路径
        question: 用户自然语言问题
        with_chart: 是否自动生成图表 (默认 True)

    Returns:
        {"answer": str, "code": str, "result": dict, "chart": dict|null}
    """
    # 1. 加载数据 (带缓存)
    safe_path = validate_file_path(file_path)
    if not os.path.exists(safe_path):
        return {"answer": f"文件不存在: {safe_path}", "code": "", "result": None, "chart": None}

    try:
        df = _get_dataframe(safe_path)
    except Exception as e:
        return {"answer": f"文件读取失败: {e}", "code": "", "result": None, "chart": None}

    # 2. 构建提示
    df_info = _build_df_info(df)
    system_prompt = DATA_CHAT_SYSTEM.format(df_info=df_info)

    # 3. LLM 调用 (不使用 with_structured_output，DeepSeek 不支持)
    try:
        llm = _get_llm()
        json_system = system_prompt + (
            "\n\n**你必须只输出一个 JSON 对象，不要包含任何其他文字、markdown 代码块标记或解释。**"
            "\nJSON 格式：{\"content\":\"markdown自然语言回答\", \"code\":\"可选的pandas代码或空字符串\", \"chart_config\":{...}或null}"
        )
        raw = llm.invoke([
            SystemMessage(content=json_system),
            HumanMessage(content=question),
        ])
        text = raw.content.strip() if hasattr(raw, 'content') else str(raw).strip()
        # 提取 JSON（处理可能的 markdown 代码块包裹）
        parsed = _extract_json(text)
        response = DataAnalysisOutput(**parsed)
    except Exception as e:
        logger.error(f"LLM 结构化输出失败: {e}")
        # 兜底：尝试修复常见的 y 为 list 问题后重试一次
        try:
            parsed_fixed = _extract_json(text)
            cc = parsed_fixed.get("chart_config") or {}
            if isinstance(cc.get("y"), list):
                y_list = cc["y"]
                cc["y"] = str(y_list[0]) if y_list else ""
                if len(y_list) > 1 and not cc.get("series"):
                    cc["series"] = [
                        {"dataKey": str(c), "chartType": "bar" if i == 0 else "line"}
                        for i, c in enumerate(y_list)
                    ]
                cc["type"] = "composed"
                parsed_fixed["chart_config"] = cc
            response = DataAnalysisOutput(**parsed_fixed)
            logger.info("LLM 输出修复成功（y 为 list → composed）")
        except Exception as e2:
            logger.error(f"修复后仍然失败: {e2}")
            return {"answer": f"LLM 调用失败: {e}", "code": "", "result": None, "chart": None}

    # 4. 执行代码
    answer = response.content
    # 兜底：过滤 LLM 可能误输出的 markdown 表格
    # 匹配: header_row \n separator_row \n data_rows
    answer = re.sub(
        r'\n\|[^\n]+\|\s*\n\|[-\s|:]+\|\s*\n(?:\|[^\n]+\|\s*\n)*',
        '\n', answer
    )
    # 清除残留的行数提示
    answer = re.sub(r'\n\*显示前\d+行.*\*', '', answer)
    code = response.code or ""
    result = None
    if code:
        try:
            result = _execute_sandbox(code, df)
        except ValueError as e:
            result = {"type": "error", "error": f"安全检查: {e}"}

    # 5. 补全答案 — 仅补充关键结果数字，表格/图表由前端结构化渲染
    if result and result.get("type") != "error":
        if result["type"] == "scalar":
            val = result.get("value")
            if val is not None and len(answer) < 120:
                answer = f"{answer}\n\n> 结果: {val}"
        elif result["type"] == "dataframe":
            # 不注入表格到回答文字 — 前端通过 data_result 事件单独渲染
            cols = result.get("columns", [])
            rows = result.get("rows", [])
            shape = result.get("shape", [0, 0])
            if cols and rows:
                preview = ", ".join(
                    f"{cols[0]}={rows[0][0]}" if rows else ""
                )[:80]
                answer += f"\n\n> 📊 数据: {shape[0]}行×{shape[1]}列，含表格和图表见下方"
        elif result["type"] == "series":
            data = result.get("data", {})
            if data:
                keys = list(data.keys())
                preview = ", ".join(f"{k}={v}" for k, v in list(data.items())[:3])
                answer += f"\n\n> 📊 {len(data)}项数据: {preview}..."

    # 6. 图表
    chart = None
    if with_chart:
        # 优先从实际执行结果生成图表数据
        if result and result.get("type") != "error":
            chart = _smart_chart_type(df, result, question)
        # LLM 给的 chart_config 中有类型/坐标轴偏好，覆盖智能推断
        if chart and response.chart_config:
            cc = response.chart_config
            # 获取 chart.data 中实际存在的 key 集合，用于校验覆盖是否安全
            data_keys = set()
            if chart.get("data") and isinstance(chart["data"], list) and chart["data"]:
                data_keys = set(chart["data"][0].keys()) if isinstance(chart["data"][0], dict) else set()
            if cc.type:
                chart["type"] = cc.type
            # 仅当 LLM 给的 x/y 在 data key 中实际存在时才覆盖，否则保留智能推断值
            if cc.x and (not data_keys or cc.x in data_keys):
                chart["x"] = cc.x
            if cc.y and (not data_keys or cc.y in data_keys):
                chart["y"] = cc.y
            if cc.x2 and (not data_keys or cc.x2 in data_keys):
                chart["x2"] = cc.x2
            if cc.title and cc.title != "图表":
                chart["title"] = cc.title
            if cc.series:
                # 校验 series 中 dataKey 是否存在于 data
                valid_series = [
                    s for s in cc.series
                    if not data_keys or s.get("dataKey") in data_keys
                ]
                if valid_series:
                    chart["series"] = valid_series
        elif not chart and response.chart_config and response.chart_config.data:
            chart = response.chart_config.model_dump()
        # 清理图表数据 (NaN/None → 0)
        chart = _sanitize_chart_data(chart)

    # 7. 数据洞察
    insights = None
    try:
        insights = _generate_data_insights(df, result)
    except Exception as e:
        logger.debug(f"数据洞察生成失败: {e}")

    # 8. 后续建议
    suggested = getattr(response, 'suggested_questions', []) or []

    return {
        "answer": answer,
        "code": code,
        "result": result,
        "chart": chart,
        "insights": insights,
        "suggested_questions": suggested,
    }


# ====== Word 报告生成 ======

def generate_data_report(
    file_path: str,
    conversation_history: list[dict] | None = None,
) -> str:
    """
    生成 Word 数据分析报告。

    Args:
        file_path: Excel/CSV 数据文件路径
        conversation_history: 多轮对话历史 [{"question": "...", "answer": "..."}, ...]

    Returns:
        报告文件路径 (data/reports/report_{timestamp}.docx)
    """
    import os as _os
    from app.tools.registry import ensure_directory

    safe_path = validate_file_path(file_path)
    if not _os.path.exists(safe_path):
        raise FileNotFoundError(f"文件不存在: {safe_path}")

    # 1. 加载数据
    ext = _os.path.splitext(safe_path)[1].lower()
    try:
        if ext in (".xlsx", ".xls"):
            df = pd.read_excel(safe_path)
            df.columns = [str(c).strip() for c in df.columns]
        elif ext == ".csv":
            for enc in ["utf-8", "gbk", "gb2312", "latin-1"]:
                try:
                    df = pd.read_csv(safe_path, encoding=enc)
                    break
                except (UnicodeDecodeError, UnicodeError):
                    continue
            else:
                raise ValueError("无法识别 CSV 编码")
        else:
            raise ValueError(f"不支持的文件格式: {ext}")
    except Exception as e:
        raise ValueError(f"文件读取失败: {e}") from e

    # 2. 生成报告概述文字 (LLM)
    df_info = _build_df_info(df)
    report_prompt = f"""你是一个数据分析师。为以下数据写一份报告摘要（200-300字）。

数据概况:
{df_info}

请生成一段专业的报告文字，包含：
1. 数据整体概况（行数列数，覆盖范围）
2. 2-3个关键发现或洞察
3. 简要的业务建议

只输出报告文字，不要JSON格式。"""

    overview_text = ""
    try:
        llm = _get_llm()
        raw = llm.invoke([HumanMessage(content=report_prompt)])
        overview_text = raw.content.strip() if hasattr(raw, 'content') else str(raw).strip()
    except Exception as e:
        logger.warning(f"LLM 报告概述生成失败: {e}")
        overview_text = f"数据包含 {len(df)} 行、{len(df.columns)} 列。"

    # 3. 生成图表 (matplotlib PNG)
    chart_paths: list[str] = []
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        output_dir = "data/reports"
        ensure_directory(output_dir)

        # 只为数值列生成图表
        numeric_cols = df.select_dtypes(include=["number"]).columns.tolist()
        categorical_cols = df.select_dtypes(include=["object"]).columns.tolist()

        if numeric_cols and categorical_cols:
            # 分类汇总柱状图
            cat_col = categorical_cols[0]
            top_cats = df.groupby(cat_col)[numeric_cols[0]].sum().nlargest(8)
            fig, ax = plt.subplots(figsize=(8, 5))
            top_cats.plot(kind="bar", ax=ax, color="#4F46E5")
            ax.set_title(f"{numeric_cols[0]} 按{cat_col}汇总 (Top 8)")
            ax.set_xlabel(cat_col)
            ax.set_ylabel(numeric_cols[0])
            plt.xticks(rotation=30, ha="right")
            fig.tight_layout()
            chart_path = _os.path.join(
                output_dir, f"chart_bar_{pd.Timestamp.now().strftime('%H%M%S_%f')}.png"
            )
            fig.savefig(chart_path, dpi=150)
            plt.close(fig)
            chart_paths.append(chart_path)

        if len(numeric_cols) >= 1:
            # 数值分布直方图
            fig, ax = plt.subplots(figsize=(8, 4))
            df[numeric_cols[0]].dropna().hist(bins=20, ax=ax, color="#10B981")
            ax.set_title(f"{numeric_cols[0]} 分布")
            ax.set_xlabel(numeric_cols[0])
            ax.set_ylabel("频次")
            fig.tight_layout()
            chart_path = _os.path.join(
                output_dir, f"chart_hist_{pd.Timestamp.now().strftime('%H%M%S_%f')}.png"
            )
            fig.savefig(chart_path, dpi=150)
            plt.close(fig)
            chart_paths.append(chart_path)
    except Exception as e:
        logger.warning(f"图表生成失败: {e}")

    # 4. 构建 Word 文档
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError:
        raise ImportError("python-docx 未安装，请运行: pip install python-docx")

    doc = Document()

    # 封面
    doc.add_heading("数据分析报告", level=0)
    doc.add_paragraph(f"生成时间: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}")
    file_name = _os.path.basename(safe_path)
    doc.add_paragraph(f"数据来源: {file_name}")
    doc.add_paragraph("")

    # 报告概述
    doc.add_heading("报告概述", level=1)
    doc.add_paragraph(overview_text)

    # 数据概览
    doc.add_heading("数据概览", level=1)
    doc.add_paragraph(f"行数: {len(df)} | 列数: {len(df.columns)}")
    doc.add_paragraph(f"列名: {', '.join(str(c) for c in df.columns)}")

    # 描述统计
    doc.add_heading("描述统计", level=1)
    desc = df.describe()
    # 构建描述统计表格
    table = doc.add_table(rows=len(desc) + 1, cols=len(desc.columns) + 1, style="Light Grid Accent 1")
    table.cell(0, 0).text = "统计量"
    for j, col in enumerate(desc.columns):
        table.cell(0, j + 1).text = str(col)
    for i, idx in enumerate(desc.index):
        table.cell(i + 1, 0).text = str(idx)
        for j, col in enumerate(desc.columns):
            val = desc.loc[idx, col]
            if isinstance(val, (int, float, np.integer, np.floating)):
                table.cell(i + 1, j + 1).text = f"{float(val):.2f}"
            else:
                table.cell(i + 1, j + 1).text = str(val)

    # 对话历史
    if conversation_history:
        doc.add_heading("分析对话", level=1)
        for i, turn in enumerate(conversation_history, 1):
            doc.add_heading(f"问题 {i}", level=2)
            doc.add_paragraph(turn.get("question", ""))
            doc.add_paragraph(turn.get("answer", ""))

    # 图表
    if chart_paths:
        doc.add_heading("可视化图表", level=1)
        for path in chart_paths:
            if _os.path.exists(path):
                try:
                    doc.add_picture(path, width=Inches(5.5))
                    caption = _os.path.basename(path).replace(".png", "")
                    doc.add_paragraph(caption).italic = True
                except Exception as e:
                    logger.warning(f"插入图片失败 {path}: {e}")

    # 保存
    ensure_directory("data/reports")
    report_path = _os.path.join(
        "data/reports",
        f"report_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S_%f')}.docx",
    )
    doc.save(report_path)
    logger.info(f"Word 报告已生成: {report_path}")
    return report_path


# ====== LangChain 工具包装 ======

class DataChatInput(BaseModel):
    file_path: str = Field(description="Excel/CSV 文件路径")
    question: str = Field(default="", description="自然语言数据分析问题")
    query: str = Field(default="", description="问题的别名 (兼容 LLM 生成)")


@register_tool
class DataConversationTool(BaseTool):
    """LLM 驱动的自然语言数据分析对话"""
    name: str = "data_conversation"
    description: str = (
        "上传数据文件后，用自然语言进行数据分析对话。"
        "支持：统计、筛选、排序、分组、趋势分析、图表生成。"
    )
    args_schema: type[BaseModel] = DataChatInput

    def _run(self, file_path: str, question: str = "", query: str = "") -> str:
        q = question or query or ""
        result = analyze_with_llm(file_path, q)
        answer = result["answer"]

        if result.get("result") and result["result"].get("type") == "error":
            answer += f"\n\n⚠️ {result['result']['error']}"
        elif result.get("result") and result["result"].get("type") == "scalar":
            val_str = str(result["result"]["value"])
            if 'base64' in val_str and len(val_str) > 500:
                val_str = val_str[:200] + "\n... [图片数据已自动省略]"
            elif len(val_str) > 500:
                val_str = val_str[:500] + "..."
            answer += f"\n\n**结果:** {val_str}"

        # 附加后续建议
        suggested = result.get("suggested_questions", [])
        if suggested:
            answer += "\n\n💡 **推荐后续分析:**\n"
            for sq in suggested[:3]:
                answer += f"- {sq}\n"

        return answer
