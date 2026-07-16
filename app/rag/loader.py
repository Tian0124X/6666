"""通用文档加载器 — PDF/Word/Excel/TXT 四种格式，自动检测文件类型"""

import os
import logging
import re
from pathlib import Path
from typing import List
from langchain_core.documents import Document

from app.rag.mineru_loader import MinerUPDFExtractor, resolve_pdf_engine, PDFEngine

logger = logging.getLogger(__name__)

# 文件头魔术数 → 真实类型（防止扩展名伪装）
FILE_SIGNATURES = {
    b"%PDF": "pdf",
    b"PK\x03\x04": "zip_based",   # docx/xlsx 都是 ZIP 格式
    b"\xd0\xcf\x11\xe0": "ole",   # 旧版 .doc/.xls
}


def detect_file_type(file_path: str) -> str:
    """通过文件头魔术数 + 扩展名双重检测"""
    ext = Path(file_path).suffix.lower()
    try:
        with open(file_path, "rb") as f:
            header = f.read(8)
        for magic, ftype in FILE_SIGNATURES.items():
            if header.startswith(magic):
                if ftype == "zip_based":
                    if ext in (".docx", ".docm"):
                        return "docx"
                    elif ext in (".pptx", ".pptm"):
                        return "pptx"
                    else:
                        return "xlsx"
                elif ftype == "ole":
                    if ext in (".doc",):
                        return "docx"
                    elif ext in (".xls",):
                        return "xlsx"
                    return "xlsx"
                return ftype
    except Exception:
        pass
    return ext.lstrip(".")


class UniversalDocumentLoader:
    """通用文档加载器，按文件类型自动分发解析器"""

    SUPPORTED_TYPES = ["pdf", "docx", "xlsx", "xls", "txt", "csv"]

    @staticmethod
    def load_pdf(file_path: str, pdf_engine: str = "auto") -> List[Document]:
        """
        PDF 解析 — 按页拆分。

        可选引擎:
          - "auto": MinerU → PyPDF2 回退
          - "mineru": 强制 MinerU
          - "pypdf2": 强制 PyPDF2

        Metadata 增强: 增加 parser 字段记录实际使用的引擎。
        """
        engine = resolve_pdf_engine(pdf_engine)

        if engine == "mineru":
            docs = UniversalDocumentLoader._load_pdf_mineru(file_path)
            # MinerU 失败(异常或无输出) -> 若是 auto 模式则回退 PyPDF2
            if not docs:
                if pdf_engine == "auto":
                    logger.warning(
                        f"MinerU 解析无输出, 回退 PyPDF2: {Path(file_path).name}"
                    )
                    return UniversalDocumentLoader._load_pdf_pypdf2(file_path)
                else:
                    logger.error(
                        f"MinerU 解析无输出 (forced engine=mineru): {file_path}"
                    )
            elif pdf_engine == "auto":
                docs = UniversalDocumentLoader._enrich_mineru_with_pypdf(file_path, docs)
            return docs
        else:
            return UniversalDocumentLoader._load_pdf_pypdf2(file_path)

    @staticmethod
    def _load_pdf_mineru(file_path: str) -> List[Document]:
        """使用 MinerU 解析 PDF，输出 Markdown 格式的按页文档。"""
        try:
            pages = MinerUPDFExtractor.extract(file_path, ocr=False)
        except Exception as e:
            logger.error(f"MinerU 解析失败: {file_path}: {e}")
            return []

        docs = []
        for page in pages:
            text = page["markdown"]
            if text and text.strip():
                meta = {
                    "source": os.path.abspath(file_path),
                    "filename": Path(file_path).name,
                    "page": page["page_number"],
                    "file_type": "pdf",
                    "parser": "mineru",
                }
                mineru_meta = page.get("metadata", {})
                if mineru_meta.get("table_count"):
                    meta["table_count"] = mineru_meta["table_count"]
                if mineru_meta.get("image_count"):
                    meta["image_count"] = mineru_meta["image_count"]
                docs.append(Document(page_content=text, metadata=meta))

        logger.info(
            f"PDF(MinerU): {file_path} → {len(docs)} 页 "
            f"(总 {sum(len(d.page_content) for d in docs)} 字符)"
        )
        return docs

    @staticmethod
    def _enrich_mineru_with_pypdf(file_path: str, mineru_docs: List[Document]) -> List[Document]:
        """仅在 PyPDF 明确提取到更多数值事实时补全文本，防止表格行静默丢失。"""
        fallback_docs = UniversalDocumentLoader._load_pdf_pypdf2(file_path)
        fallback_by_page = {item.metadata.get("page"): item for item in fallback_docs}
        enriched: list[Document] = []
        for document in mineru_docs:
            fallback = fallback_by_page.get(document.metadata.get("page"))
            if not fallback:
                enriched.append(document)
                continue
            mineru_numbers = set(re.findall(r"\d+(?:\.\d+)?", document.page_content))
            fallback_numbers = set(re.findall(r"\d+(?:\.\d+)?", fallback.page_content))
            missing_numbers = fallback_numbers - mineru_numbers
            mineru_table_rows = set(re.findall(r"已满\s*\d+\s*年(?:\s*不满\s*\d+\s*年)?", document.page_content))
            fallback_table_rows = set(re.findall(r"已满\s*\d+\s*年(?:\s*不满\s*\d+\s*年)?", fallback.page_content))
            missing_table_rows = fallback_table_rows - mineru_table_rows
            if len(missing_numbers) < 3 and len(missing_table_rows) < 2:
                enriched.append(document)
                continue
            metadata = dict(document.metadata)
            metadata["parser"] = f"{metadata.get('parser', 'mineru')}+pypdf_fallback"
            metadata["fallback_numeric_facts"] = len(missing_numbers)
            metadata["fallback_table_rows"] = len(missing_table_rows)
            enriched.append(Document(
                page_content=document.page_content.rstrip() + "\n\n[PyPDF 表格文本回退]\n" + fallback.page_content,
                metadata=metadata,
            ))
        return enriched

    @staticmethod
    def _load_pdf_pypdf2(file_path: str) -> List[Document]:
        """使用 pypdf / PyPDF2 解析 PDF，自动回退。

        优先使用 pypdf（PyPDF2 的维护版，CJK 支持更好），
        不可用时回退 PyPDF2。
        """
        # 优先使用 pypdf（PyPDF2 的继任者，中文提取更好）
        PdfReader = None
        engine_name = "pypdf2"
        for lib_name, import_path in [
            ("pypdf", "pypdf"),
            ("PyPDF2", "PyPDF2"),
        ]:
            try:
                mod = __import__(import_path, fromlist=["PdfReader"])
                PdfReader = mod.PdfReader
                engine_name = lib_name
                break
            except ImportError:
                continue

        if PdfReader is None:
            logger.error(f"PDF 解析失败: 未安装 pypdf 或 PyPDF2，请执行 pip install pypdf")
            return []

        reader = PdfReader(file_path)
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception:
                logger.warning(f"PDF 加密且无法解密: {file_path}")
                return []

        docs = []
        for i, page in enumerate(reader.pages):
            try:
                text = page.extract_text()
            except Exception:
                text = ""
            if text and text.strip():
                docs.append(Document(
                    page_content=text.strip(),
                    metadata={
                        "source": os.path.abspath(file_path),
                        "filename": Path(file_path).name,
                        "page": i + 1,
                        "file_type": "pdf",
                        "parser": engine_name,
                    },
                ))

        if not docs:
            logger.warning(
                f"PDF ({engine_name}) 未提取到文本: {file_path} "
                f"(共 {len(reader.pages)} 页) — 可能是扫描件/图片型 PDF，"
                f"建议安装 MinerU 进行 OCR 解析: pip install magic-pdf"
            )

        logger.info(f"PDF ({engine_name}): {file_path} → {len(docs)} 页")
        return docs

    @staticmethod
    def load_docx(file_path: str) -> List[Document]:
        """Word 解析"""
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if not full_text:
            return []
        logger.info(f"DOCX: {file_path} → {len(full_text)} 字符")
        return [Document(
            page_content=full_text,
            metadata={
                "source": os.path.abspath(file_path),
                "filename": Path(file_path).name,
                "file_type": "docx",
            },
        )]

    @staticmethod
    def load_excel(file_path: str) -> List[Document]:
        """Excel 解析 — 每行转'列名: 值' 格式，方便语义检索"""
        import pandas as pd
        sheets = pd.read_excel(file_path, sheet_name=None)
        docs = []
        for sheet_name, df in sheets.items():
            rows_text = []
            for _, row in df.iterrows():
                row_str = " | ".join(
                    f"{col}: {val}" for col, val in row.items()
                    if pd.notna(val)
                )
                if row_str.strip():
                    rows_text.append(row_str)
            if not rows_text:
                continue
            content = f"[Sheet: {sheet_name}]\n" + "\n".join(rows_text)
            docs.append(Document(
                page_content=content,
                metadata={
                    "source": os.path.abspath(file_path),
                    "filename": Path(file_path).name,
                    "sheet": sheet_name,
                    "rows": len(df),
                    "file_type": "excel",
                },
            ))
        logger.info(f"Excel: {file_path} → {len(docs)} sheets")
        return docs

    @staticmethod
    def load_txt(file_path: str) -> List[Document]:
        """TXT/CSV 解析 — 自动检测编码"""
        for enc in ["utf-8", "gbk", "gb2312", "latin-1"]:
            try:
                with open(file_path, "r", encoding=enc) as f:
                    text = f.read()
                if text.strip():
                    logger.info(f"TXT: {file_path} → {len(text)} 字符 (编码: {enc})")
                    return [Document(
                        page_content=text,
                        metadata={
                            "source": os.path.abspath(file_path),
                            "filename": Path(file_path).name,
                            "encoding": enc,
                            "file_type": "txt",
                        },
                    )]
            except (UnicodeDecodeError, UnicodeError):
                continue
        logger.warning(f"TXT: {file_path} 编码检测失败")
        return []

    @classmethod
    def load(cls, file_path: str, pdf_engine: str = "auto") -> List[Document]:
        """主入口 — 自动检测类型并加载

        Args:
            file_path: 文件路径
            pdf_engine: PDF 解析引擎 ("auto"/"mineru"/"pypdf2")

        Returns:
            List[Document]
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        file_type = detect_file_type(file_path)

        loaders = {
            "pdf": cls.load_pdf,
            "docx": cls.load_docx,
            "xlsx": cls.load_excel,
            "xls": cls.load_excel,
            "txt": cls.load_txt,
            "csv": cls.load_txt,
        }

        loader = loaders.get(file_type)
        if loader is None:
            raise ValueError(
                f"不支持的文件类型: {file_type}。支持 {cls.SUPPORTED_TYPES}"
            )

        # Only pass pdf_engine to PDF loader; other loader methods don't accept it
        if file_type == "pdf":
            return loader(file_path, pdf_engine=pdf_engine)
        return loader(file_path)
